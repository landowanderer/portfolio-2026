#!/usr/bin/env python3
"""Build a Google Docs-ready DOCX from Harry's verified job research."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
JOBS_PATH = ROOT / "data" / "jobs.json"
OUTPUT_PATH = (
    ROOT / "output" / "Harry_Career_Opportunities_2026-07-26_Google_Docs.docx"
)
VERIFY_DATE = "2026-07-26"
EAST_ASIA_FONT = "Arial Unicode MS"

PRIMARY_IDS = [
    "skims-4430168363",
    "lagom-4442693326",
    "jacques-marie-mage-121407",
    "alpinestars-9d277ad6",
    "aninebing-5282740008",
    "haus-motion-78902CC55A",
    "directive-57dc2d91",
    "melin-4417018055",
]

SECONDARY_IDS = [
    "quilt-5186941007",
    "helpscout-82c4c13c",
    "insomniac-4306642171",
    "umg-26991",
    "hinge-health-beb00f5c",
    "nen-creative-a7c735e1",
    "cotton-citizen-hDtw4giwXz",
    "triumph-5811befc",
]

CONTRACT_IDS = [
    "webtoon-6fd68ef9",
    "intro-59e8c47c",
    "hybe-4443070657",
    "handshake-1c09f77c",
]

STRETCH_IDS = [
    "la28-6564235003",
    "fabletics-R8203",
    "trueshort-9efb1967",
]


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = EAST_ASIA_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Arial")
    font.set(qn("w:hAnsi"), "Arial")
    font.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    properties.extend([font, color, underline])
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_label_paragraph(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    label_run = paragraph.add_run(f"{label}：")
    set_run_font(label_run, 10.5, True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, 10.5)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.16)
    paragraph.paragraph_format.space_after = Pt(2)
    for run in paragraph.runs:
        set_run_font(run, 10.5)
    if not paragraph.runs:
        set_run_font(paragraph.add_run(text), 10.5)


def salary_text(job: dict) -> str:
    minimum = job["salary_min"]
    maximum = job["salary_max"]
    if minimum is None:
        return "未注明"
    currency = job["salary_currency"] or "USD"
    if currency == "USD/hour":
        if minimum == maximum:
            return f"${minimum:,.0f}/小时"
        return f"${minimum:,.0f}–${maximum:,.0f}/小时"
    return f"${minimum:,.0f}–${maximum:,.0f}/年"


def status_text(job: dict) -> str:
    return f"可申请；官方职位页/ATS 于 {job['last_verified']} 核验"


def authorization_summary(job: dict) -> str:
    sponsorship = job["sponsorship_text"]
    work_auth = job["work_authorization_text"]
    if "No visa sponsorship" in sponsorship or "do not sponsor" in sponsorship:
        return f"{sponsorship} 投递前请确认与你最新 OPT 计划是否兼容。"
    if "STEM OPT is not supported" in sponsorship:
        return (
            f"{work_auth} {sponsorship} 接受任何工作前需与学校 DSO 确认；本文不作移民法律判断。"
        )
    if work_auth.startswith("Unknown") and sponsorship.startswith("Unknown"):
        return "职位未说明；若表单询问，必须由你根据最新情况亲自回答。"
    return f"{work_auth} {sponsorship}"


def add_job(document: Document, job: dict, index: int) -> None:
    heading = document.add_paragraph(style="Heading 3")
    heading.paragraph_format.keep_with_next = True
    number_run = heading.add_run(f"{index}. ")
    set_run_font(number_run, 14, True)
    title_run = heading.add_run(f"{job['company']} — {job['title']}")
    set_run_font(title_run, 14, True)

    add_label_paragraph(
        document,
        "基本信息",
        (
            f"{job['location']}｜{job['remote_type']}｜{job['employment_type']}｜"
            f"{salary_text(job)}｜匹配分 {job['fit_score']}/100"
        ),
    )
    add_label_paragraph(document, "开放状态", status_text(job))
    add_label_paragraph(document, "经验要求", job["experience_required"])
    add_label_paragraph(document, "身份/签证提醒", authorization_summary(job))

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    lead = paragraph.add_run("为什么值得投：")
    set_run_font(lead, 10.5, True)
    set_run_font(paragraph.add_run(job["fit_reasons"][0]), 10.5)
    add_bullet(document, job["fit_reasons"][1])

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    lead = paragraph.add_run("主要缺口：")
    set_run_font(lead, 10.5, True)
    set_run_font(paragraph.add_run(job["gaps"][0]), 10.5)
    add_bullet(document, job["gaps"][1])

    add_label_paragraph(
        document,
        "作品集顺序",
        " → ".join(job["recommended_projects"]),
    )
    add_label_paragraph(document, "下一步", job["next_action"])

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(10)
    lead = paragraph.add_run("官方入口：")
    set_run_font(lead, 10.5, True)
    add_hyperlink(paragraph, "打开职位并自己申请", job["canonical_url"])
    linkedin_urls = [
        url
        for url in job.get("also_seen_on", [])
        if "linkedin.com/jobs/" in url
    ]
    if linkedin_urls:
        set_run_font(paragraph.add_run("　｜　"), 10.5)
        add_hyperlink(paragraph, "LinkedIn 备用链接", linkedin_urls[0])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = EAST_ASIA_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(6)

    heading_specs = {
        "Heading 1": (20, 10, 6),
        "Heading 2": (16, 8, 4),
        "Heading 3": (14, 7, 3),
    }
    for name, (size, before, after) in heading_specs.items():
        style = document.styles[name]
        style.font.name = EAST_ASIA_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Small Note" not in [style.name for style in document.styles]:
        small_note = document.styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
        small_note.base_style = document.styles["Normal"]
        small_note.font.name = EAST_ASIA_FONT
        small_note._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        small_note.font.size = Pt(9)
        small_note.font.color.rgb = RGBColor(80, 80, 80)
        small_note.paragraph_format.space_after = Pt(4)


def add_section(
    document: Document,
    title: str,
    intro: str,
    ids: list[str],
    jobs_by_id: dict[str, dict],
    start_index: int,
    page_break: bool = True,
) -> int:
    if page_break:
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    paragraph = document.add_paragraph(intro)
    for run in paragraph.runs:
        set_run_font(run, 11)
    index = start_index
    for job_id in ids:
        add_job(document, jobs_by_id[job_id], index)
        index += 1
    return index


def build() -> Path:
    payload = json.loads(JOBS_PATH.read_text(encoding="utf-8"))
    jobs_by_id = {job["job_id"]: job for job in payload["jobs"]}
    required_ids = PRIMARY_IDS + SECONDARY_IDS + CONTRACT_IDS + STRETCH_IDS
    missing = [job_id for job_id in required_ids if job_id not in jobs_by_id]
    if missing:
        raise RuntimeError(f"Missing jobs: {', '.join(missing)}")

    document = Document()
    configure_document(document)
    document.core_properties.title = "Harry 求职岗位清单"
    document.core_properties.subject = "已核实仍可申请的设计与创意科技岗位"
    document.core_properties.author = "Harry Career Agent"

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("Harry 求职岗位清单")
    set_run_font(title_run, 26, False)

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run("已核实仍可申请｜更新于 2026 年 7 月 26 日")
    set_run_font(subtitle_run, 14, True)
    subtitle.paragraph_format.space_after = Pt(16)

    summary = document.add_paragraph()
    summary_run = summary.add_run(
        "本清单包含 23 个在 2026 年 7 月 26 日仍显示职位详情与申请入口的岗位，"
        "其中 7 个为本轮新增。优先使用公司官方申请页；只有未找到独立公司页面时才使用"
        "已现场核验的 LinkedIn 职位链接。"
    )
    set_run_font(summary_run, 11)

    document.add_heading("怎么使用这份清单", level=1)
    instructions = [
        "先看“本地/远程优先”，挑 2–3 个今天申请；不要一次性海投。",
        "点击每个岗位末尾的“打开职位并自己申请”，再核对职位标题和申请表是否仍正常显示。",
        "工作授权、未来是否需要 sponsorship、薪资期望等敏感问题必须由你本人按最新情况回答。",
        "任何岗位都没有被自动投递；这份文档只做研究、排序和申请准备。",
    ]
    for item in instructions:
        add_bullet(document, item)

    document.add_heading("今天建议先处理", level=2)
    priorities = [
        "SKIMS — Jr Graphic Designer：最近毕业生到 1 年经验；电商、动效、视频与 AI 工具高度匹配。",
        "LAGOM — Junior Graphic Designer：1–3 年；品牌、短视频、包装、动效与生成式 AI 都相关。",
        "Jacques Marie Mage — Junior Graphic Designer：洛杉矶、1–2 年、品牌/摄影/包装方向。",
        "Alpinestars — Junior Graphic Designer：Torrance、1–2 年，实习与自由职业计入。",
        "ANINE BING — Associate Graphic Designer：洛杉矶，时尚品牌与多渠道视觉执行。",
        "melin — Jr. Graphic Designer：1–2 年，但每周四天在 San Clemente，先判断通勤现实性。",
    ]
    for item in priorities:
        add_bullet(document, item)

    document.add_heading("核验时排除的旧岗位", level=2)
    exclusions = [
        "goodr — Associate Designer：旧 requisition 现跳转至公司职位总表并显示 error=true；今天移除。",
        "Zip 与 Centerfield 的旧初级设计岗位仍不在官方当前职位列表。",
        "Kikoff、Foxglove 与 Funko 的旧设计岗位仍不可申请。",
        "UMG 旧 Santa Monica requisition 已关闭；清单只保留当前 Philadelphia 版本。",
    ]
    for item in exclusions:
        add_bullet(document, item)

    index = 1
    index = add_section(
        document,
        "A. 本地 / 远程优先",
        "这些岗位在地点、年资或核心能力上最接近 Harry。建议先投 SKIMS、LAGOM、Jacques Marie Mage、Alpinestars 与 ANINE BING。",
        PRIMARY_IDS,
        jobs_by_id,
        index,
    )
    index = add_section(
        document,
        "B. 第二梯队与搬迁选项",
        "这些岗位仍然可以申请，但存在搬迁、通勤、无 sponsorship、行业专项技能或已上线产品证据等更明显的缺口。",
        SECONDARY_IDS,
        jobs_by_id,
        index,
        page_break=False,
    )
    index = add_section(
        document,
        "C. 合同 / 自由职业",
        "这些项目与普通全职岗位分开管理。接受任何合同或临时职位前，要确认期限、工时、雇佣分类、OPT/DSO 要求和税务责任。",
        CONTRACT_IDS,
        jobs_by_id,
        index,
        page_break=False,
    )
    add_section(
        document,
        "D. 选择性冲刺",
        "LA28 与 Fabletics 的内容匹配，但年资要求偏高；TrueShort 的创意科技方向相关，但技术门槛也明显高于目前材料能证明的能力。",
        STRETCH_IDS,
        jobs_by_id,
        index,
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("投递前的统一检查", level=1)
    checklist = [
        "职位页面仍显示 Apply / Submit application。",
        "简历版本与岗位路线一致：Visual/Brand、Product/Experience 或 Creative Technology。",
        "作品集首屏与文档建议顺序一致，且 Anaago 仅在已获准公开时使用。",
        "不填写未经确认的工作授权、sponsorship、薪资、残障或人口统计信息。",
        "不把概念项目写成已上线产品，不虚构客户、结果、指标、工具或工作年限。",
        "提交后记录日期、使用的简历版本、作品集顺序和下一次 follow-up 日期。",
    ]
    for item in checklist:
        add_bullet(document, item)

    note = document.add_paragraph(style="Small Note")
    note.add_run(
        "说明：匹配分只用于排序，不代表录取概率；移民/工作授权内容仅摘录职位原文，"
        "不构成法律意见。岗位状态会变化，正式提交前仍需由你再次打开官方链接确认。"
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build())
